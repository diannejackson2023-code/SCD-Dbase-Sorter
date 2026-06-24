import docx
import pandas as pd
import os
import sys
# Handle both relative and absolute imports
if __name__ != '__main__':
    try:
        from .mapping import find_master_match, MASTER_HEADINGS
    except ImportError:
        from mapping import find_master_match, MASTER_HEADINGS
else:
    from mapping import find_master_match, MASTER_HEADINGS

def process_word_file(file_path):
    """
    Parses a Word file (.docx) for SCD records.
    Scans tables for recognized headers and extracts row data.
    """
    doc = docx.Document(file_path)
    all_data = []
    
    for table in doc.tables:
        if len(table.rows) < 2:
            continue
            
        # Determine column mapping from the first row
        header_row = [cell.text.strip() for cell in table.rows[0].cells]
        mapping = {}
        for i, text in enumerate(header_row):
            match = find_master_match(text)
            if match:
                mapping[i] = match
        
        # If no recognized headers, try the second row (handles merged title rows)
        if not mapping and len(table.rows) > 2:
            header_row = [cell.text.strip() for cell in table.rows[1].cells]
            for i, text in enumerate(header_row):
                match = find_master_match(text)
                if match:
                    mapping[i] = match
                    
        # Extract data if mapping was found
        if mapping:
            start_row = 1 if 0 in mapping or any(m in mapping.values() for m in MASTER_HEADINGS) else 2
            for row in table.rows[start_row:]:
                row_data = {}
                for i, master_name in mapping.items():
                    row_data[master_name] = row.cells[i].text.strip()
                if any(row_data.values()): # Only add non-empty rows
                    all_data.append(row_data)
                    
    return pd.DataFrame(all_data)

def scan_text_for_keywords(file_path, keywords=None):
    """
    Scans a Word file for specific keywords to confirm it's an SCD-related document.
    """
    if keywords is None:
        keywords = ["scd", "sickle cell", "hemoglobin", "genotype", "hydroxyurea"]
        
    doc = docx.Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text.lower())
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text.lower())
                
    content = " ".join(full_text)
    found_keywords = [k for k in keywords if k in content]
    return len(found_keywords) > 0, found_keywords
